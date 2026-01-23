#!/usr/bin/env python3
"""
NRG Legislative Intelligence - V2 Full Pipeline

End-to-end pipeline that:
1. Fetches bills from configured sources (Open States, Congress.gov, etc.)
2. Runs Sequential Evolution analysis on all versions
3. Applies Two-Tier Validation (Judge + Multi-sample + Fallback)
4. Scores findings using rubric-based assessment
5. Generates reports in multiple formats (JSON, Markdown, DOCX)

Usage:
    python run_v2_full_pipeline.py [options]

Options:
    --verbose, -v       Enable verbose logging (show all traces)
    --quiet, -q         Quiet mode (minimal output)
    --config FILE       Path to config file (default: config.yaml)
    --output DIR        Output directory for reports (default: ./reports)

Examples:
    python run_v2_full_pipeline.py                    # Normal mode
    python run_v2_full_pipeline.py -v                 # Verbose mode
    python run_v2_full_pipeline.py --output ./out     # Custom output dir
"""
import os
import sys
import json
import time
import tempfile
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from enum import Enum

import yaml
import httpx
from dotenv import load_dotenv
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.tree import Tree
from rich import box

# Import V2 analysis components
from nrg_core.v2.two_tier import TwoTierOrchestrator
from nrg_core.v2.sequential_evolution import SequentialEvolutionAgent, BillVersion
from nrg_core.config import load_nrg_context

load_dotenv()

# ============================================================================
# LOGGING SYSTEM
# ============================================================================

class LogLevel(Enum):
    QUIET = 0
    NORMAL = 1
    VERBOSE = 2
    DEBUG = 3


@dataclass
class PipelineLogger:
    """Centralized logger with trace support for pipeline visibility."""
    console: Console = field(default_factory=Console)
    level: LogLevel = LogLevel.NORMAL
    trace_config: dict = field(default_factory=dict)
    show_timestamps: bool = True
    show_token_counts: bool = True
    show_cost_estimates: bool = True

    def _format_time(self) -> str:
        if self.show_timestamps and self.level.value >= LogLevel.VERBOSE.value:
            return f"[dim]{datetime.now().strftime('%H:%M:%S')}[/dim] "
        return ""

    def info(self, message: str, emoji: str = ""):
        """Standard info message."""
        if self.level.value >= LogLevel.NORMAL.value:
            prefix = f"{emoji} " if emoji else ""
            self.console.print(f"{self._format_time()}{prefix}{message}")

    def verbose(self, message: str, emoji: str = ""):
        """Verbose message (only in verbose/debug mode)."""
        if self.level.value >= LogLevel.VERBOSE.value:
            prefix = f"{emoji} " if emoji else ""
            self.console.print(f"{self._format_time()}[dim]{prefix}{message}[/dim]")

    def debug(self, message: str):
        """Debug message (only in debug mode)."""
        if self.level.value >= LogLevel.DEBUG.value:
            self.console.print(f"{self._format_time()}[dim cyan]DEBUG: {message}[/dim cyan]")

    def success(self, message: str):
        """Success message."""
        if self.level.value >= LogLevel.NORMAL.value:
            self.console.print(f"{self._format_time()}[green]✓[/green] {message}")

    def warning(self, message: str):
        """Warning message."""
        self.console.print(f"{self._format_time()}[yellow]⚠[/yellow] {message}")

    def error(self, message: str):
        """Error message."""
        self.console.print(f"{self._format_time()}[red]✗[/red] {message}")

    def stage_header(self, stage_num: int, title: str, description: str = ""):
        """Log a pipeline stage header."""
        if self.level.value >= LogLevel.NORMAL.value:
            self.console.print()
            self.console.rule(f"[bold cyan]Stage {stage_num}: {title}[/bold cyan]", style="cyan")
            if description and self.level.value >= LogLevel.VERBOSE.value:
                self.console.print(f"[dim]{description}[/dim]")

    def trace_decision(self, component: str, decision: str, rationale: str, details: dict = None):
        """Log a decision with rationale - core for traceability."""
        trace_enabled = self.trace_config.get(component, False)
        if self.level.value >= LogLevel.VERBOSE.value and trace_enabled:
            self.console.print(f"  [bold yellow]→ Decision:[/bold yellow] {decision}")
            self.console.print(f"    [dim]Rationale: {rationale}[/dim]")
            if details and self.level.value >= LogLevel.DEBUG.value:
                for k, v in details.items():
                    self.console.print(f"    [dim]{k}: {v}[/dim]")

    def trace_finding(self, finding_num: int, statement: str, quotes: list, confidence: float, impact: int):
        """Log finding extraction details."""
        if self.level.value >= LogLevel.VERBOSE.value and self.trace_config.get("sequential_evolution", False):
            tree = Tree(f"[bold]Finding F{finding_num}[/bold]")
            tree.add(f"Statement: {statement[:100]}{'...' if len(statement) > 100 else ''}")
            if quotes:
                quotes_branch = tree.add("Supporting Quotes")
                for q in quotes[:2]:
                    text = q.get("text", "") if isinstance(q, dict) else str(q)
                    section = q.get("section", "?") if isinstance(q, dict) else "?"
                    quotes_branch.add(f'[dim]"{text[:60]}..." (§{section})[/dim]')
            tree.add(f"Confidence: [{'green' if confidence >= 0.8 else 'yellow' if confidence >= 0.6 else 'red'}]{confidence:.2f}[/]")
            tree.add(f"Impact Estimate: [{'red' if impact >= 7 else 'yellow' if impact >= 4 else 'green'}]{impact}/10[/]")
            self.console.print(tree)

    def trace_validation(self, finding_num: int, quote_verified: bool, hallucination: bool,
                         evidence_quality: float, judge_confidence: float, reason: str = ""):
        """Log validation result with explanation."""
        if self.level.value >= LogLevel.VERBOSE.value and self.trace_config.get("two_tier_validation", False):
            status_icon = "[green]✓[/green]" if quote_verified and not hallucination else "[red]✗[/red]"
            status_text = "Verified" if quote_verified and not hallucination else "Hallucination" if hallucination else "Unverified"

            self.console.print(f"  F{finding_num}: {status_icon} {status_text}")
            self.console.print(f"    [dim]Quote Verified: {quote_verified} | Evidence Quality: {evidence_quality:.2f} | Judge Confidence: {judge_confidence:.2f}[/dim]")

            if hallucination and self.trace_config.get("hallucination_detection", False):
                self.console.print(f"    [yellow]→ Hallucination Reason: {reason or 'Finding claims impact not supported by bill text'}[/yellow]")

    def trace_rubric_score(self, finding_id: str, dimension: str, score: int, rationale: str, anchor: str):
        """Log rubric scoring details."""
        if self.level.value >= LogLevel.VERBOSE.value and self.trace_config.get("rubric_scoring", False):
            color = "green" if score <= 2 else "yellow" if score <= 5 else "red" if score <= 8 else "bold red"
            self.console.print(f"  {dimension}: [{color}]{score}/10[/{color}] → {anchor}")
            if self.level.value >= LogLevel.DEBUG.value and rationale:
                self.console.print(f"    [dim]Rationale: {rationale[:150]}...[/dim]")

    def trace_stability(self, finding_id: str, score: float, origin_version: int, mod_count: int, prediction: str):
        """Log stability analysis."""
        if self.level.value >= LogLevel.VERBOSE.value and self.trace_config.get("stability_analysis", False):
            color = "green" if score >= 0.8 else "yellow" if score >= 0.5 else "red"
            self.console.print(f"  {finding_id}: Stability [{color}]{score:.2f}[/{color}]")
            self.console.print(f"    [dim]Origin: v{origin_version} | Modifications: {mod_count} | Prediction: {prediction}[/dim]")

    def trace_complexity(self, bill_id: str, score: int, route: str, reasons: list):
        """Log complexity assessment."""
        if self.level.value >= LogLevel.VERBOSE.value:
            self.console.print(f"  [bold]Complexity Assessment:[/bold] {bill_id}")
            self.console.print(f"    Score: {score} → Route: [cyan]{route}[/cyan]")
            for reason in reasons:
                self.console.print(f"    [dim]• {reason}[/dim]")

    def show_table(self, title: str, columns: list, rows: list):
        """Display a formatted table."""
        if self.level.value >= LogLevel.NORMAL.value:
            table = Table(title=title, box=box.ROUNDED)
            for col in columns:
                table.add_column(col)
            for row in rows:
                table.add_row(*[str(x) for x in row])
            self.console.print(table)


# ============================================================================
# CONFIGURATION
# ============================================================================

def load_config(config_path: str = "config.yaml") -> dict:
    """Load configuration from YAML file."""
    try:
        with open(config_path, "r") as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        print(f"Error: {config_path} not found")
        sys.exit(1)


def get_logger_from_config(config: dict, cli_level: Optional[str] = None) -> PipelineLogger:
    """Create logger from config."""
    log_config = config.get("logging", {})

    # CLI overrides config
    if cli_level == "verbose":
        level = LogLevel.VERBOSE
    elif cli_level == "quiet":
        level = LogLevel.QUIET
    elif cli_level == "debug":
        level = LogLevel.DEBUG
    else:
        level_str = log_config.get("level", "normal")
        level = LogLevel[level_str.upper()] if level_str.upper() in LogLevel.__members__ else LogLevel.NORMAL

    trace_config = log_config.get("trace", {
        "sequential_evolution": True,
        "two_tier_validation": True,
        "rubric_scoring": True,
        "hallucination_detection": True,
        "stability_analysis": True,
    })
    format_config = log_config.get("format", {})

    return PipelineLogger(
        level=level,
        trace_config=trace_config,
        show_timestamps=format_config.get("show_timestamps", True),
        show_token_counts=format_config.get("show_token_counts", True),
        show_cost_estimates=format_config.get("show_cost_estimates", True),
    )


# ============================================================================
# PDF TEXT EXTRACTION
# ============================================================================

def extract_pdf_text(pdf_url: str, logger: PipelineLogger) -> str:
    """Extract text from PDF URL using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed - skipping PDF extraction")
        return ""

    try:
        logger.verbose(f"Downloading PDF from {pdf_url[:60]}...")
        with httpx.Client(timeout=30.0, follow_redirects=True) as http:
            response = http.get(pdf_url)
            response.raise_for_status()

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(response.content)
                tmp_path = tmp.name

            text_parts = []
            with pdfplumber.open(tmp_path) as pdf:
                logger.verbose(f"Extracting text from {len(pdf.pages)} pages...")
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            os.unlink(tmp_path)
            full_text = "\n".join(text_parts)
            logger.verbose(f"Extracted {len(full_text)} characters ({len(full_text.split())} words)")
            return full_text

    except Exception as e:
        logger.debug(f"PDF extraction failed: {e}")
        return ""


# ============================================================================
# DATA FETCHING
# ============================================================================

def fetch_openstates_bills(jurisdiction: str, bill_numbers: List[str], logger: PipelineLogger) -> List[Dict]:
    """Fetch bills from Open States API."""
    api_key = os.getenv("OPENSTATES_API_KEY")
    if not api_key:
        logger.warning("OPENSTATES_API_KEY not set")
        return []

    logger.info(f"Fetching {len(bill_numbers)} bills from Open States ({jurisdiction})...")

    base_url = "https://v3.openstates.org"
    headers = {
        "X-API-KEY": api_key,
        "User-Agent": "NRG-Energy-Legislative-Tracker/2.0"
    }
    bills = []

    try:
        with httpx.Client(timeout=60.0) as http:
            for bill_num in bill_numbers:
                logger.verbose(f"Fetching {jurisdiction} {bill_num}...")

                # Search for bill
                search_url = f"{base_url}/bills"
                params = {
                    "jurisdiction": jurisdiction.lower(),
                    "q": bill_num,
                    "per_page": 5
                }

                response = http.get(search_url, headers=headers, params=params)
                response.raise_for_status()
                data = response.json()

                for result in data.get("results", []):
                    identifier = result.get("identifier", "")
                    if identifier.replace(" ", "").upper() == bill_num.replace(" ", "").upper():
                        openstates_id = result.get("id", "")

                        # Fetch versions
                        versions = []
                        if openstates_id:
                            detail_url = f"{base_url}/bills/{openstates_id}"
                            detail_params = {"include": "versions"}
                            detail_resp = http.get(detail_url, headers=headers, params=detail_params)
                            if detail_resp.status_code == 200:
                                detail_data = detail_resp.json()
                                raw_versions = detail_data.get("versions", [])

                                for idx, v in enumerate(raw_versions, 1):
                                    version_text = ""
                                    links = v.get("links", [])

                                    # Try to get PDF text
                                    for link in links:
                                        if link.get("media_type") == "application/pdf":
                                            pdf_url = link.get("url")
                                            version_text = extract_pdf_text(pdf_url, logger)
                                            if version_text:
                                                break

                                    if version_text:
                                        versions.append({
                                            "version_number": idx,
                                            "version_type": v.get("note", f"Version {idx}"),
                                            "version_date": v.get("date", ""),
                                            "full_text": version_text,
                                            "word_count": len(version_text.split())
                                        })

                        # Get latest version text for primary analysis
                        bill_text = ""
                        if versions:
                            bill_text = versions[-1].get("full_text", "")

                        if bill_text:
                            bills.append({
                                "source": "Open States",
                                "type": f"{jurisdiction} State Bill",
                                "bill_id": identifier.replace(" ", ""),
                                "number": identifier,
                                "title": result.get("title", "No title"),
                                "status": result.get("actions", [{}])[-1].get("description", "Unknown") if result.get("actions") else "Unknown",
                                "bill_text": bill_text,
                                "versions": versions,
                                "openstates_id": openstates_id,
                                "jurisdiction": jurisdiction
                            })
                            logger.success(f"Found {identifier} ({len(versions)} versions)")
                        break

                time.sleep(0.3)  # Rate limiting

    except Exception as e:
        logger.error(f"Error fetching from Open States: {e}")

    logger.success(f"Fetched {len(bills)} bills from Open States")
    return bills


def fetch_bills_from_config(config: dict, logger: PipelineLogger) -> List[Dict]:
    """Fetch all bills configured in sources."""
    bills = []

    # Open States Texas bills
    openstates_config = config.get("sources", {}).get("openstates", {})
    if openstates_config.get("enabled"):
        texas_config = openstates_config.get("texas_bills", {})
        if texas_config.get("enabled"):
            jurisdiction = texas_config.get("jurisdiction", "TX")
            bill_numbers = texas_config.get("bills", [])
            if bill_numbers:
                texas_bills = fetch_openstates_bills(jurisdiction, bill_numbers, logger)
                bills.extend(texas_bills)

    return bills


# ============================================================================
# COMPLEXITY ASSESSMENT
# ============================================================================

def assess_complexity(bill: dict, config: dict, logger: PipelineLogger) -> tuple:
    """Assess bill complexity and determine routing."""
    v2_config = config.get("v2", {}).get("orchestration", {})
    scoring = v2_config.get("scoring", {
        "pages_20_50": 1,
        "pages_over_50": 2,
        "versions_2_5": 1,
        "versions_over_5": 2,
        "domain_environmental": 1,
        "domain_energy_tax": 2,
    })
    thresholds = v2_config.get("complexity_thresholds", {"standard_max": 2})

    score = 0
    reasons = []

    # Version count
    version_count = len(bill.get("versions", []))
    if version_count > 5:
        score += scoring.get("versions_over_5", 2)
        reasons.append(f"Many versions ({version_count}) +{scoring.get('versions_over_5', 2)}")
    elif version_count >= 2:
        score += scoring.get("versions_2_5", 1)
        reasons.append(f"Multiple versions ({version_count}) +{scoring.get('versions_2_5', 1)}")

    # Page count estimate
    total_text = sum(len(v.get("full_text", "")) for v in bill.get("versions", []))
    page_estimate = total_text // 3000
    if page_estimate > 50:
        score += scoring.get("pages_over_50", 2)
        reasons.append(f"Long bill (~{page_estimate} pages) +{scoring.get('pages_over_50', 2)}")
    elif page_estimate >= 20:
        score += scoring.get("pages_20_50", 1)
        reasons.append(f"Medium length (~{page_estimate} pages) +{scoring.get('pages_20_50', 1)}")

    # Domain detection
    title_lower = bill.get("title", "").lower()
    if any(kw in title_lower for kw in ["energy", "tax", "oil", "gas", "power"]):
        score += scoring.get("domain_energy_tax", 2)
        reasons.append(f"Energy/Tax domain +{scoring.get('domain_energy_tax', 2)}")
    elif any(kw in title_lower for kw in ["environment", "emission", "climate", "pollution"]):
        score += scoring.get("domain_environmental", 1)
        reasons.append(f"Environmental domain +{scoring.get('domain_environmental', 1)}")

    # Determine route
    standard_max = thresholds.get("standard_max", 2)
    route = "STANDARD" if score <= standard_max else "ENHANCED"

    logger.trace_complexity(bill.get("bill_id", "Unknown"), score, route, reasons)

    return route, score, reasons


# ============================================================================
# REPORT GENERATION
# ============================================================================

def generate_markdown_report(results: dict, output_path: Path, logger: PipelineLogger):
    """Generate detailed Markdown report."""
    logger.verbose("Generating Markdown report...")

    lines = [
        "# NRG Legislative Intelligence Report",
        "",
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**Pipeline:** V2 (Sequential Evolution + Two-Tier Validation)",
        f"**Provider:** {results.get('provider', 'Unknown')}",
        f"**Bills Analyzed:** {results.get('bills_analyzed', 0)}",
        f"**Total Time:** {results.get('elapsed_seconds', 0):.1f}s",
        "",
        "---",
        "",
    ]

    for bill_result in results.get("results", []):
        bill_id = bill_result.get("bill_id", "Unknown")
        title = bill_result.get("bill_title", "No title")

        lines.extend([
            f"## {bill_id}: {title}",
            "",
            f"| Metric | Value |",
            f"|--------|-------|",
            f"| Source | {bill_result.get('source', 'Unknown')} |",
            f"| Route | {bill_result.get('route', 'STANDARD')} |",
            f"| Versions Processed | {bill_result.get('versions_processed', 0)} |",
            f"| Total Findings | {bill_result.get('findings_count', 0)} |",
            f"| Verified Findings | {bill_result.get('verified_findings', 0)} |",
            f"| Hallucinations | {bill_result.get('hallucinations_detected', 0)} |",
            "",
            "### Findings",
            "",
        ])

        for i, finding in enumerate(bill_result.get("findings", []), 1):
            validations = bill_result.get("validations", [])
            validation = validations[i - 1] if i <= len(validations) else {}
            verified = "✓" if validation.get("quote_verified") and not validation.get("hallucination_detected") else "✗"
            halluc_tag = " [Hallucination]" if validation.get("hallucination_detected") else ""

            lines.extend([
                f"#### Finding {i} {verified}{halluc_tag}",
                "",
                f"**Statement:** {finding.get('statement', 'No statement')}",
                "",
                f"- Confidence: {finding.get('confidence', 0):.2f}",
                f"- Impact Estimate: {finding.get('impact_estimate', 0)}/10",
                "",
            ])

            if finding.get("quotes"):
                lines.append("**Quotes:**")
                for quote in finding.get("quotes", []):
                    lines.append(f"> \"{quote.get('text', '')}\" — §{quote.get('section', '?')}")
                lines.append("")

        # Rubric scores summary
        rubric_scores = bill_result.get("rubric_scores", [])
        if rubric_scores:
            lines.extend([
                "### Rubric Scores",
                "",
                "| Dimension | Score | Anchor |",
                "|-----------|-------|--------|",
            ])
            seen = set()
            for score in rubric_scores:
                dim = score.get("dimension", "")
                if dim not in seen:
                    seen.add(dim)
                    lines.append(f"| {dim} | {score.get('score', 0)}/10 | {score.get('rubric_anchor', '')} |")
            lines.append("")

        lines.extend(["", "---", ""])

    output_path.write_text("\n".join(lines))
    logger.success(f"Markdown report: {output_path}")


def generate_docx_report(results: dict, output_path: Path, logger: PipelineLogger):
    """Generate DOCX report."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed, skipping DOCX")
        return

    logger.verbose("Generating DOCX report...")

    doc = Document()
    title = doc.add_heading("NRG Legislative Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Pipeline: V2 (Sequential Evolution + Two-Tier Validation)")
    doc.add_paragraph(f"Provider: {results.get('provider', 'Unknown')}")
    doc.add_paragraph(f"Bills Analyzed: {results.get('bills_analyzed', 0)}")

    for bill_result in results.get("results", []):
        doc.add_heading(f"{bill_result.get('bill_id', 'Unknown')}: {bill_result.get('bill_title', '')}", level=1)

        # Summary table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"

        metrics = [
            ("Source", bill_result.get("source", "")),
            ("Route", bill_result.get("route", "")),
            ("Versions", str(bill_result.get("versions_processed", 0))),
            ("Findings", str(bill_result.get("findings_count", 0))),
            ("Verified", str(bill_result.get("verified_findings", 0))),
            ("Hallucinations", str(bill_result.get("hallucinations_detected", 0))),
        ]
        for metric, value in metrics:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = value

        # Findings
        doc.add_heading("Findings", level=2)
        for i, finding in enumerate(bill_result.get("findings", []), 1):
            validations = bill_result.get("validations", [])
            validation = validations[i - 1] if i <= len(validations) else {}
            status = "✓" if validation.get("quote_verified") and not validation.get("hallucination_detected") else "✗"

            doc.add_heading(f"Finding {i} {status}", level=3)
            doc.add_paragraph(finding.get("statement", ""))
            doc.add_paragraph(f"Confidence: {finding.get('confidence', 0):.2f} | Impact: {finding.get('impact_estimate', 0)}/10")

    doc.save(str(output_path))
    logger.success(f"DOCX report: {output_path}")


def generate_reports(results: dict, output_dir: Path, config: dict, logger: PipelineLogger) -> Path:
    """Generate all configured report formats."""
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"nrg_analysis_{timestamp}"

    formats = config.get("output", {}).get("formats", {"json": True, "markdown": True, "docx": True})

    # JSON (always)
    json_path = output_dir / f"{base_name}.json"
    with open(json_path, "w") as f:
        json.dump(results, f, indent=2)
    logger.success(f"JSON report: {json_path}")

    # Markdown
    if formats.get("markdown", True):
        generate_markdown_report(results, output_dir / f"{base_name}.md", logger)

    # DOCX
    if formats.get("docx", True):
        generate_docx_report(results, output_dir / f"{base_name}.docx", logger)

    return json_path


# ============================================================================
# MAIN ANALYSIS PIPELINE
# ============================================================================

def analyze_bill_v2(bill: Dict[str, Any], config: dict, nrg_context: str, api_key: str, logger: PipelineLogger) -> Dict[str, Any]:
    """Run full V2 analysis pipeline on a single bill."""
    bill_id = bill.get("bill_id", bill.get("number", "Unknown"))
    bill_text = bill.get("bill_text", "")

    if not bill_text:
        logger.error(f"No bill text available for {bill_id}")
        return {"error": "No bill text available", "bill_id": bill_id}

    # Stage 0: Complexity Assessment
    route, complexity_score, complexity_reasons = assess_complexity(bill, config, logger)

    # Stage 1: Sequential Evolution
    logger.stage_header(1, "Sequential Evolution", "Walk versions chronologically, extract findings, track modifications")

    # Convert versions to BillVersion objects
    versions = []
    raw_versions = bill.get("versions", [])
    if raw_versions:
        for v in raw_versions:
            if v.get("full_text"):
                versions.append(BillVersion(
                    version_number=v["version_number"],
                    text=v["full_text"],
                    name=v.get("version_type", f"Version {v['version_number']}")
                ))
    else:
        versions = [BillVersion(version_number=1, text=bill_text, name="Current")]

    logger.verbose(f"Processing {len(versions)} versions sequentially...")

    try:
        evolution_agent = SequentialEvolutionAgent(api_key=api_key)
        evolution_result = evolution_agent.walk_versions(bill_id=bill_id, versions=versions)

        findings_count = len(evolution_result.findings_registry)
        logger.success(f"Extracted {findings_count} findings from {len(versions)} versions")

        # Log each finding with details
        for i, finding in enumerate(evolution_result.findings_registry, 1):
            quotes = [{"text": q.text, "section": q.section} for q in finding.quotes] if hasattr(finding, 'quotes') else []
            logger.trace_finding(
                i,
                finding.statement if hasattr(finding, 'statement') else str(finding),
                quotes,
                finding.confidence if hasattr(finding, 'confidence') else 0.5,
                finding.impact_estimate if hasattr(finding, 'impact_estimate') else 0
            )

            # Log stability for this finding
            if hasattr(finding, 'origin_version') and hasattr(finding, 'modification_count'):
                stability = evolution_result.stability_scores.get(finding.id, 0.5) if hasattr(finding, 'id') else 0.5
                prediction = "Very stable" if stability >= 0.9 else "Stable" if stability >= 0.7 else "May change" if stability >= 0.5 else "Volatile"
                logger.trace_stability(
                    f"F{i}",
                    stability,
                    finding.origin_version,
                    finding.modification_count,
                    prediction
                )

    except Exception as e:
        logger.error(f"Sequential Evolution failed: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e), "bill_id": bill_id}

    # Stage 2: Two-Tier Validation
    logger.stage_header(2, "Two-Tier Validation", "Validate findings, detect hallucinations, assign confidence")

    try:
        orchestrator = TwoTierOrchestrator(
            judge_api_key=api_key,
            enable_multi_sample=config.get("v2", {}).get("two_tier", {}).get("multi_sample", {}).get("enabled", True),
            enable_fallback=config.get("v2", {}).get("two_tier", {}).get("fallback", {}).get("enabled", True)
        )

        result = orchestrator.validate(
            bill_id=bill_id,
            bill_text=bill_text,
            nrg_context=nrg_context,
            findings_registry=evolution_result.findings_registry,
            stability_scores=evolution_result.stability_scores
        )

        # Log validation for each finding
        for i, validation in enumerate(result.judge_validations, 1):
            reason = ""
            if validation.hallucination_detected:
                reason = "Claim about NRG impact not grounded in bill text"

            logger.trace_validation(
                i,
                validation.quote_verified,
                validation.hallucination_detected,
                validation.evidence_quality,
                validation.judge_confidence,
                reason
            )

        verified_count = sum(1 for v in result.judge_validations if v.quote_verified and not v.hallucination_detected)
        halluc_count = sum(1 for v in result.judge_validations if v.hallucination_detected)
        logger.success(f"Validated: {verified_count} verified, {halluc_count} hallucinations filtered")

        # Log decision about multi-sample or fallback
        if result.multi_sample_agreement is not None:
            logger.trace_decision(
                "two_tier_validation",
                f"Multi-sample check: {result.multi_sample_agreement:.1%} agreement",
                f"Ran {config.get('v2', {}).get('two_tier', {}).get('multi_sample', {}).get('samples', 3)} samples to verify consistency"
            )

        if result.second_model_reviewed:
            logger.trace_decision(
                "two_tier_validation",
                "Fallback model invoked",
                "Judge confidence was uncertain (0.6-0.8) and impact >= 7"
            )

    except Exception as e:
        logger.error(f"Two-Tier Validation failed: {e}")
        import traceback
        traceback.print_exc()
        return {"error": str(e), "bill_id": bill_id}

    # Stage 3: Rubric Scoring
    logger.stage_header(3, "Rubric Scoring", "Score validated findings on 4 dimensions: legal, financial, operational, ambiguity")

    rubric_scores = result.rubric_scores
    if rubric_scores:
        logger.verbose(f"Scoring {len(rubric_scores)} dimension assessments...")

        # Group and log by finding
        current_finding = None
        for score in rubric_scores:
            logger.trace_rubric_score(
                "",
                score.dimension,
                score.score,
                score.rationale,
                score.rubric_anchor
            )

        logger.success(f"Generated {len(rubric_scores)} rubric scores")

    # Compile result
    analysis_result = {
        "bill_id": bill_id,
        "bill_title": bill.get("title", ""),
        "source": bill.get("source", ""),
        "route": route,
        "complexity_score": complexity_score,
        "versions_processed": len(versions),
        "findings_count": len(result.primary_analysis.findings),
        "findings": [
            {
                "statement": f.statement,
                "quotes": [{"text": q.text, "section": q.section} for q in f.quotes],
                "confidence": f.confidence,
                "impact_estimate": f.impact_estimate
            }
            for f in result.primary_analysis.findings
        ],
        "validations": [
            {
                "finding_id": v.finding_id,
                "quote_verified": v.quote_verified,
                "hallucination_detected": v.hallucination_detected,
                "evidence_quality": v.evidence_quality,
                "judge_confidence": v.judge_confidence
            }
            for v in result.judge_validations
        ],
        "rubric_scores": [
            {
                "dimension": s.dimension,
                "score": s.score,
                "rationale": s.rationale[:200] + "..." if len(s.rationale) > 200 else s.rationale,
                "rubric_anchor": s.rubric_anchor
            }
            for s in rubric_scores
        ],
        "multi_sample_agreement": result.multi_sample_agreement,
        "second_model_reviewed": result.second_model_reviewed,
        "hallucinations_detected": halluc_count,
        "verified_findings": verified_count,
    }

    # Display summary panel
    logger.console.print(Panel(
        f"[bold]{bill_id}[/bold]: {bill.get('title', '')[:60]}...\n\n"
        f"Route: {route} (complexity: {complexity_score})\n"
        f"Versions: {len(versions)}\n\n"
        f"[bold]Findings:[/bold] {len(result.primary_analysis.findings)} total\n"
        f"  [green]Verified:[/green] {verified_count}\n"
        f"  [red]Hallucinations:[/red] {halluc_count}\n"
        f"  Rubric Scores: {len(rubric_scores)}",
        title=f"Analysis Complete: {bill_id}",
        border_style="green" if halluc_count == 0 else "yellow" if halluc_count < 3 else "red",
    ))

    return analysis_result


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """Main entry point for V2 pipeline."""
    parser = argparse.ArgumentParser(
        description="NRG Legislative Intelligence V2 Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("-v", "--verbose", action="store_true", help="Verbose logging")
    parser.add_argument("-q", "--quiet", action="store_true", help="Quiet mode")
    parser.add_argument("--debug", action="store_true", help="Debug logging")
    parser.add_argument("--config", default="config.yaml", help="Config file path")
    parser.add_argument("--output", help="Output directory for reports")
    args = parser.parse_args()

    # Load configuration
    config = load_config(args.config)

    # Set up logger
    cli_level = "debug" if args.debug else "verbose" if args.verbose else "quiet" if args.quiet else None
    logger = get_logger_from_config(config, cli_level)

    # Determine output directory
    output_dir = Path(args.output) if args.output else Path(config.get("output", {}).get("directory", "./reports"))

    # Get API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logger.error("OPENAI_API_KEY not set")
        sys.exit(1)

    # Load NRG context
    nrg_context = load_nrg_context()

    # Show startup banner
    provider = config.get("llm", {}).get("provider", "openai")
    logger.console.print(Panel(
        f"[bold cyan]NRG Legislative Intelligence v2.0[/bold cyan]\n"
        f"Provider: {provider}\n"
        f"Architecture: Sequential Evolution + Two-Tier Validation\n"
        f"Logging: {logger.level.name}\n"
        f"Output: {output_dir}",
        title="Pipeline Starting",
        border_style="cyan",
    ))

    start_time = time.time()

    # Fetch bills
    bills = fetch_bills_from_config(config, logger)
    if not bills:
        logger.error("No bills fetched - check config.yaml and API keys")
        sys.exit(1)

    # Analyze each bill
    logger.info(f"Analyzing {len(bills)} bills with V2 pipeline...")
    results = []

    for idx, bill in enumerate(bills, 1):
        bill_id = bill.get("bill_id", "Unknown")
        logger.console.print()
        logger.console.rule(f"[bold]({idx}/{len(bills)}) {bill_id}[/bold]")

        analysis = analyze_bill_v2(bill, config, nrg_context, api_key, logger)
        results.append(analysis)

    elapsed = time.time() - start_time

    # Compile final output
    output = {
        "pipeline": "v2",
        "provider": provider,
        "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
        "elapsed_seconds": elapsed,
        "bills_analyzed": len(results),
        "results": results,
    }

    # Generate reports
    logger.console.print()
    logger.console.rule("[bold]Generating Reports[/bold]")
    report_path = generate_reports(output, output_dir, config, logger)

    # Show final summary
    total_findings = sum(r.get("findings_count", 0) for r in results)
    total_verified = sum(r.get("verified_findings", 0) for r in results)
    total_halluc = sum(r.get("hallucinations_detected", 0) for r in results)

    logger.console.print()
    logger.console.print(Panel(
        f"[bold green]Analysis Complete[/bold green]\n\n"
        f"Bills Analyzed: {len(results)}\n"
        f"Total Findings: {total_findings}\n"
        f"Verified: {total_verified}\n"
        f"Hallucinations Filtered: {total_halluc}\n"
        f"Total Time: {elapsed:.1f}s\n\n"
        f"Reports saved to: {output_dir}",
        title="Summary",
        border_style="green",
    ))

    return 0


if __name__ == "__main__":
    sys.exit(main())
